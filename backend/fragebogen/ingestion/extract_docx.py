"""Strukturelle Extraktion aus docx (Spec §6.2).

Zwei Quellen:
1. Tabellen: typisches Fragebogen-Layout ist eine Tabelle mit Spalten
   „Frage/Question" + „Antwort/Answer". Jede Datenzeile → eine Frage; die
   Antwort-Zelle (Tabelle/Zeile/Spalte) wird als ``feld_referenz`` gemerkt.
2. Absätze: Fließtext-Fragen werden anhand eines Fragezeichens bzw. einer
   führenden Nummerierung erkannt.
"""

from __future__ import annotations

import re

_FRAGE_KEYS = ("frage", "question")
_ANTWORT_KEYS = ("antwort", "answer", "kommentar", "comment", "response")
_NUMMER_PREFIX = re.compile(r"^\s*\d+[.)]\s*")


def extrahiere_fragen_docx(pfad: str) -> list[dict]:
    """Extrahiert Fragen aus Tabellen + Absätzen; kaputte Datei → leere Liste."""
    from docx import Document

    try:
        doc = Document(pfad)
    except Exception:
        return []

    fragen: list[dict] = []
    fragen.extend(_aus_tabellen(doc))
    fragen.extend(_aus_absaetzen(doc))
    return fragen


def _aus_tabellen(doc) -> list[dict]:
    fragen: list[dict] = []
    for t_idx, tabelle in enumerate(doc.tables):
        if not tabelle.rows:
            continue
        kopf = [_clean(z.text) for z in tabelle.rows[0].cells]
        frage_col = _spalte_mit_key(kopf, _FRAGE_KEYS)
        antwort_col = _spalte_mit_key(kopf, _ANTWORT_KEYS)
        hat_header = frage_col is not None
        if not hat_header:
            frage_col = 0
            antwort_col = 1 if len(kopf) > 1 else None
        startzeile = 1 if hat_header else 0
        for r_idx, zeile in enumerate(tabelle.rows[startzeile:], start=startzeile):
            zellen = zeile.cells
            text = _clean(zellen[frage_col].text) if frage_col < len(zellen) else ""
            if not text:
                continue
            ref = {"tabelle": t_idx, "zeile": r_idx, "spalte": frage_col}
            if antwort_col is not None and antwort_col < len(zellen):
                ref["antwort_spalte"] = antwort_col
            fragen.append({"nummer": None, "text": text, "feld_referenz": ref})
    return fragen


def _aus_absaetzen(doc) -> list[dict]:
    fragen: list[dict] = []
    for p_idx, absatz in enumerate(doc.paragraphs):
        text = _clean(absatz.text)
        if not _ist_frage(text):
            continue
        nummer_match = _NUMMER_PREFIX.match(text)
        nummer = nummer_match.group().strip(" .)") if nummer_match else None
        fragen.append(
            {
                "nummer": nummer,
                "text": text,
                "feld_referenz": {"absatz": p_idx},
            }
        )
    return fragen


def _ist_frage(text: str) -> bool:
    if not text:
        return False
    return text.endswith("?") or bool(_NUMMER_PREFIX.match(text))


def _spalte_mit_key(kopf: list[str], keys: tuple[str, ...]) -> int | None:
    for col, zelle in enumerate(kopf):
        low = zelle.lower()
        if any(k in low for k in keys):
            return col
    return None


def _clean(text: str) -> str:
    return (text or "").strip()
