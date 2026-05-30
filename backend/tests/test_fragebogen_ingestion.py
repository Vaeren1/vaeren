"""Tests für die strukturelle Ingestion-Schicht (Phase D).

Fixture-Dateien werden im Test erzeugt (kein committetes Binär-Asset).
Rein strukturelle Extraktion — kein echter LLM/OCR-Call.
"""

from __future__ import annotations

import openpyxl

from fragebogen.ingestion.detect import erkenne_format_und_tier
from fragebogen.ingestion.extract_docx import extrahiere_fragen_docx
from fragebogen.ingestion.extract_pdfform import extrahiere_fragen_pdfform
from fragebogen.ingestion.extract_xlsx import extrahiere_fragen_xlsx

# --- D1: Format-/Tier-Erkennung ------------------------------------------


def test_xlsx_ist_tier1(tmp_path):
    p = tmp_path / "q.xlsx"
    wb = openpyxl.Workbook()
    wb.active["A1"] = "Frage?"
    wb.save(p)
    wb.close()
    fmt, tier = erkenne_format_und_tier(str(p))
    assert fmt == "xlsx"
    assert tier == 1


def test_pdf_acroform_ist_tier1(tmp_path):
    p = _erzeuge_fillable_pdf(tmp_path, feldname="antwort_1")
    fmt, tier = erkenne_format_und_tier(str(p))
    assert fmt == "pdf_form"
    assert tier == 1


def test_pdf_mit_text_ist_tier2(tmp_path):
    p = _erzeuge_text_pdf(tmp_path, "Haben Sie ein ISMS?")
    fmt, tier = erkenne_format_und_tier(str(p))
    assert fmt == "pdf_unstrukturiert"
    assert tier == 2


def test_pdf_ohne_text_ist_tier3(tmp_path):
    p = _erzeuge_leeres_pdf(tmp_path)
    fmt, tier = erkenne_format_und_tier(str(p))
    assert fmt == "pdf_unstrukturiert"
    assert tier == 3


def test_docx_ist_tier3_ohne_formularfelder(tmp_path):
    from docx import Document

    p = tmp_path / "d.docx"
    doc = Document()
    doc.add_paragraph("Irgendein Text")
    doc.save(p)
    fmt, tier = erkenne_format_und_tier(str(p))
    assert fmt == "docx"
    assert tier == 3


def test_unbekanntes_format_ist_tier3(tmp_path):
    p = tmp_path / "x.txt"
    p.write_text("hallo")
    fmt, tier = erkenne_format_und_tier(str(p))
    assert fmt == "txt"
    assert tier == 3


def test_kaputte_pdf_crasht_nicht(tmp_path):
    p = tmp_path / "kaputt.pdf"
    p.write_bytes(b"%PDF-1.4 das ist kein gueltiges PDF")
    # Darf nicht crashen, fällt auf Tier 3 zurück.
    fmt, tier = erkenne_format_und_tier(str(p))
    assert fmt == "pdf_unstrukturiert"
    assert tier == 3


# --- D2: xlsx-Extraktion --------------------------------------------------


def test_extract_xlsx_findet_fragen(tmp_path):
    p = tmp_path / "vda.xlsx"
    wb = openpyxl.Workbook()
    ws = wb.active
    ws["A1"] = "Nr"
    ws["B1"] = "Frage"
    ws["C1"] = "Antwort"
    ws["A2"] = "1"
    ws["B2"] = "Haben Sie ein ISMS?"
    ws["C2"] = ""
    wb.save(p)
    wb.close()
    fragen = extrahiere_fragen_xlsx(str(p))
    assert fragen[0]["text"] == "Haben Sie ein ISMS?"
    assert fragen[0]["feld_referenz"]["antwort_zelle"] == "C2"
    assert fragen[0]["feld_referenz"]["sheet"] == ws.title
    assert fragen[0]["nummer"] == "1"


def test_extract_xlsx_englischer_header(tmp_path):
    p = tmp_path / "en.xlsx"
    wb = openpyxl.Workbook()
    ws = wb.active
    ws["A1"] = "No"
    ws["B1"] = "Question"
    ws["C1"] = "Answer"
    ws["A2"] = "1"
    ws["B2"] = "Do you have an ISMS?"
    ws["A3"] = "2"
    ws["B3"] = "Do you encrypt data?"
    wb.save(p)
    wb.close()
    fragen = extrahiere_fragen_xlsx(str(p))
    assert len(fragen) == 2
    assert fragen[1]["text"] == "Do you encrypt data?"
    assert fragen[1]["feld_referenz"]["antwort_zelle"] == "C3"


def test_extract_xlsx_fallback_ohne_header(tmp_path):
    p = tmp_path / "nohead.xlsx"
    wb = openpyxl.Workbook()
    ws = wb.active
    # Keine erkennbare Header-Zeile: erste Textspalte = Frage, nächste leere = Antwort.
    ws["A1"] = "Beschreiben Sie Ihre Backup-Strategie."
    ws["A2"] = "Wie oft testen Sie Restores?"
    wb.save(p)
    wb.close()
    fragen = extrahiere_fragen_xlsx(str(p))
    assert len(fragen) == 2
    assert fragen[0]["text"] == "Beschreiben Sie Ihre Backup-Strategie."
    assert fragen[0]["feld_referenz"]["antwort_zelle"] == "B1"


def test_extract_xlsx_leere_datei(tmp_path):
    p = tmp_path / "leer.xlsx"
    wb = openpyxl.Workbook()
    wb.save(p)
    wb.close()
    fragen = extrahiere_fragen_xlsx(str(p))
    assert fragen == []


# --- D3: PDF-Formular + docx ---------------------------------------------


def test_extract_pdfform_findet_felder(tmp_path):
    p = _erzeuge_fillable_pdf(tmp_path, feldname="frage_isms")
    fragen = extrahiere_fragen_pdfform(str(p))
    assert len(fragen) == 1
    assert fragen[0]["feld_referenz"]["feldname"] == "frage_isms"


def test_extract_pdfform_kaputt_gibt_leer(tmp_path):
    p = tmp_path / "kaputt.pdf"
    p.write_bytes(b"nicht wirklich pdf")
    assert extrahiere_fragen_pdfform(str(p)) == []


def test_extract_docx_findet_tabellenfragen(tmp_path):
    from docx import Document

    p = tmp_path / "tab.docx"
    doc = Document()
    tabelle = doc.add_table(rows=3, cols=2)
    tabelle.cell(0, 0).text = "Frage"
    tabelle.cell(0, 1).text = "Antwort"
    tabelle.cell(1, 0).text = "Haben Sie ein ISMS?"
    tabelle.cell(1, 1).text = ""
    tabelle.cell(2, 0).text = "Verschlüsseln Sie Daten?"
    tabelle.cell(2, 1).text = ""
    doc.save(p)
    fragen = extrahiere_fragen_docx(str(p))
    texte = [f["text"] for f in fragen]
    assert "Haben Sie ein ISMS?" in texte
    assert "Verschlüsseln Sie Daten?" in texte


def test_extract_docx_findet_absatzfragen(tmp_path):
    from docx import Document

    p = tmp_path / "abs.docx"
    doc = Document()
    doc.add_paragraph("1. Beschreiben Sie Ihre Zugriffskontrolle?")
    doc.add_paragraph("Antwort:")
    doc.save(p)
    fragen = extrahiere_fragen_docx(str(p))
    texte = [f["text"] for f in fragen]
    assert any("Zugriffskontrolle" in t for t in texte)


def test_extract_docx_kaputt_gibt_leer(tmp_path):
    p = tmp_path / "kaputt.docx"
    p.write_bytes(b"kein docx")
    assert extrahiere_fragen_docx(str(p)) == []


# --- Fixture-Helper -------------------------------------------------------


def _erzeuge_text_pdf(tmp_path, text: str):
    from reportlab.pdfgen import canvas

    p = tmp_path / "text.pdf"
    c = canvas.Canvas(str(p))
    c.drawString(100, 750, text)
    c.save()
    return p


def _erzeuge_leeres_pdf(tmp_path):
    from reportlab.pdfgen import canvas

    p = tmp_path / "leer.pdf"
    c = canvas.Canvas(str(p))
    c.showPage()  # leere Seite ohne Text
    c.save()
    return p


def _erzeuge_fillable_pdf(tmp_path, feldname: str):
    """Minimales PDF mit einem AcroForm-Textfeld."""
    from reportlab.pdfgen import canvas

    p = tmp_path / "form.pdf"
    c = canvas.Canvas(str(p))
    c.drawString(100, 750, "Frage 1")
    c.acroForm.textfield(
        name=feldname,
        x=100,
        y=700,
        width=200,
        height=20,
        borderStyle="inset",
    )
    c.save()
    return p
